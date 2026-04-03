from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break_before = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break_before)


def apply_run_fonts(paragraph, east_asia: str, latin: str, size_pt: float, bold=False) -> None:
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = latin
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        run._element.rPr.rFonts.set(qn("w:ascii"), latin)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def apply_normal(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    apply_run_fonts(paragraph, "宋体", "Times New Roman", 10.5)


def apply_heading(paragraph, level: int) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(10)
    size = 18 if level == 1 else 15 if level == 2 else 12
    apply_run_fonts(paragraph, "黑体", "Times New Roman", size, bold=True)


def apply_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = Pt(0)
    apply_run_fonts(paragraph, "宋体", "Times New Roman", 10.5)


def apply_reference(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = Pt(-21)
    paragraph.paragraph_format.left_indent = Pt(21)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    apply_run_fonts(paragraph, "宋体", "Times New Roman", 10.5)


def load_image_map(path: Path | None) -> dict[str, Path]:
    if path is None or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    return {k: Path(v) for k, v in data.items()}


def add_image(doc: Document, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(14.5))


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for row in lines:
        cells = [c.strip().replace("`", "") for c in row.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    headers = rows[0]
    data_rows = rows[2:]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for r_idx, row in enumerate(data_rows, start=1):
        for c_idx, text in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = text
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = 1.25
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_run_fonts(p, "宋体", "Times New Roman", 10.5)


def build_doc(source: Path, image_map: dict[str, Path]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    lines = source.read_text(encoding="utf-8").splitlines()
    current_section = ""
    in_code = False
    code_lang = ""
    pending_mermaid = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if in_code:
            if stripped.startswith("```"):
                in_code = False
                if code_lang == "mermaid":
                    pending_mermaid = True
                code_lang = ""
            elif code_lang != "mermaid":
                p = doc.add_paragraph()
                p.add_run(line.rstrip())
                apply_run_fonts(p, "Consolas", "Consolas", 9)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip().lower()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(stripped[2:].strip())
            apply_run_fonts(p, "黑体", "Times New Roman", 18, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.add_run(text)
            apply_heading(p, 1)
            if text not in {"摘  要", "Abstract"}:
                add_page_break_before(p)
            current_section = text
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(stripped[4:].strip())
            apply_heading(p, 2)
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.add_run(stripped[5:].strip())
            apply_heading(p, 3)
            i += 1
            continue

        image_match = re.match(r"^\[此处插入截图：(.+?)\]$", stripped)
        if image_match:
            label = image_match.group(1).strip()
            image_path = image_map.get(label)
            if image_path and image_path.exists():
                add_image(doc, image_path)
            i += 1
            continue

        if re.match(r"^(图|表)\s*\d+(\.\d+)?", stripped):
            if pending_mermaid:
                image_path = image_map.get(stripped)
                if image_path and image_path.exists():
                    add_image(doc, image_path)
                pending_mermaid = False
            p = doc.add_paragraph()
            p.add_run(stripped.replace("`", ""))
            apply_caption(p)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        p = doc.add_paragraph()
        p.add_run(stripped.replace("`", ""))
        if current_section == "参 考 文 献" and re.match(r"^\[\d+\]", stripped):
            apply_reference(p)
        else:
            apply_normal(p)
        i += 1

    return doc


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: python generate_thesis_docx.py <source.md> <target.docx> [image-map.json]")
        return 1

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    image_map = load_image_map(Path(sys.argv[3])) if len(sys.argv) >= 4 else {}
    target.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc(source, image_map)
    doc.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
